import streamlit as st
from docx import Document
import PyPDF2
import json
import requests
from io import BytesIO
import time
from queue import Queue
from threading import Thread
from textblob import TextBlob  # For Sentiment Analysis

# Define the API endpoint and access token
API_URL = "https://api.psnext.info/api/chat"
PSCHATACCESSTOKEN = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJVc2VySW5mbyI6eyJpZCI6MzcxMzcsInJvbGVzIjpbImRlZmF1bHQiXSwicGF0aWQiOiI0MTZjZjhhYy0wZTk3LTRkMzctYWU2Yy02ZGYwMjVmMzg4MjQifSwiaWF0IjoxNzMxNDM3MDY2LCJleHAiOjE3MzQwMjkwNjZ9.QV_W6DRL488LJlQSQtDElaL50n3TOojAbitwXx15YLo"  # Replace with your actual access token

# Rate limiting variables
requests_per_minute = 5  # Set your API rate limit here
rate_limit_delay = 60 / requests_per_minute  # Time to wait between requests in seconds

# Queue to manage requests
request_queue = Queue()
results = {}

# Function to extract text from a Word document
def extract_text_from_word(docx_file):
    doc = Document(docx_file)
    return '\n'.join([para.text for para in doc.paragraphs])

# Function to extract text from a PDF file
def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    pdf_text = ""
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        pdf_text += page.extract_text() or ""  # Ensure no None values
    return pdf_text

# Function to extract text from a TXT file
def extract_text_from_txt(txt_file):
    return txt_file.read().decode('utf-8')

# Function to extract text from different file types
def extract_text_from_file(file):
    if file.type == "application/pdf":
        return extract_text_from_pdf(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_word(file)
    elif file.type == "text/plain":
        return extract_text_from_txt(file)
    else:
        return None

# Sentiment analysis function (TextBlob)
def analyze_sentiment(text):
    blob = TextBlob(text)
    sentiment_score = blob.sentiment.polarity  # Sentiment score between -1 (negative) and 1 (positive)
    return sentiment_score

# Function to evaluate interview performance (candidate, interviewer, questions, and responses)
def evaluate_interview(transcript_text, job_title):
    payload = {
    "message": (
        "You are tasked with providing a comprehensive and detailed analysis of the interview transcript. "
        f"The candidate is applying for the role of '{job_title}'. Use this context to evaluate both the interviewer and the candidate based on their relevance to this specific role.\n\n"
        
        "The analysis should address the performance of both the interviewer and the candidate, covering the following aspects:\n\n"
        
        "1. **Interviewer Evaluation**: Assess the clarity, relevance, and quality of the questions asked by the interviewer. "
        "Were the questions engaging and open-ended? Did the interviewer ask relevant follow-up questions that explored key skills and competencies specific to the role? "
        "Did they adapt their questions to the candidate’s responses? Was the interviewer’s tone appropriate (e.g., respectful, encouraging)? "
        "Did the interviewer create a comfortable and open environment for the candidate? "
        "Evaluate if the interviewer demonstrated a thorough understanding of the role and requirements and assessed the candidate's cultural fit.\n\n"
        
        "2. **Candidate Evaluation**: Evaluate the clarity, depth, and relevance of the candidate’s responses, especially in relation to the role of '{job_title}'. "
        "Did the candidate provide specific examples, skills, and accomplishments to support their answers? "
        "Was the candidate's tone confident, positive, and engaging? Were their responses structured in a clear and coherent manner? "
        "Did the candidate demonstrate a strong understanding of the role and the company? "
        "Evaluate if the candidate addressed the interview questions thoroughly and effectively conveyed their qualifications and motivations.\n\n"
        
        "3. **Sentiment Analysis**: Provide a sentiment evaluation for both the interviewer’s questions and the candidate’s responses. "
        "Was the overall tone of the interview positive, neutral, or negative? "
        "Provide a sentiment score for both the interviewer and the candidate, and indicate the general tone of the interaction. "
        "For example, was the interview collaborative or did it feel more confrontational? Were there signs of nervousness, confidence, or disengagement? "
        "Was the tone friendly, professional, or distant?\n\n"
        
        "4. **Interview Coverage**: Evaluate if the interviewer covered all the critical areas necessary for the job position. "
        "Did they focus on both technical skills and soft skills? Did they inquire about key aspects such as leadership, problem-solving, teamwork, and communication? "
        "Were there any crucial areas that were not explored, such as the candidate’s long-term career goals or cultural fit? "
        "Also, assess if the candidate addressed all relevant topics in their responses. Did they talk about past experiences, key achievements, and challenges? "
        "Did they discuss how they would approach the role and fit into the company culture? "
        "Assess if there was any lack of clarity or missed opportunities for the candidate to elaborate on important points.\n\n"
        
        "5. **Interpersonal Dynamics**: Evaluate the overall interaction between the interviewer and the candidate. "
        "How well did they build rapport with each other? Did the candidate seem at ease or uncomfortable? "
        "Was the interviewer actively listening and responding appropriately, or were they interrupting the candidate too much? "
        "Did they maintain eye contact (if in person) or create a professional atmosphere (if virtual)? "
        "Assess the effectiveness of non-verbal communication such as tone, pace, and pauses in the conversation.\n\n"
        
        "6. **Strengths and Areas for Improvement**: Highlight the strengths of both the interviewer and the candidate. "
        "For the interviewer, mention the positive aspects of their questioning, tone, and engagement. "
        "For the candidate, focus on their communication strengths, such as clarity, confidence, and relevant experiences shared. "
        "Also, provide constructive feedback on areas for improvement. For the interviewer, could they improve their questioning technique or tone? "
        "For the candidate, are there areas where they could elaborate more or refine their responses? "
        "Suggest improvements that would help both parties perform better in future interviews.\n\n"
        
        "7. **Overall Impression**: Based on the analysis, provide an overall evaluation of the interview. "
        "Was it a successful interview, and why? Did both the interviewer and candidate achieve their goals? "
        "Summarize the key takeaways and final recommendations for both the interviewer and the candidate.\n\n"
        
        f"Transcript Text:\n{transcript_text}"
    ),
    "options": {"model": "gpt35turbo"}
}


    headers = {
        "Authorization": f"Bearer {PSCHATACCESSTOKEN}",
        "Content-Type": "application/json"
    }

    response = requests.post(API_URL, headers=headers, json=payload)

    if response.status_code == 200:
        response_data = response.json()
        messages = response_data.get('data', {}).get('messages', [])
        for message in messages:
            if message.get('role') == 'assistant':
                return message.get('content', 'No content returned from the API.')
        return 'No assistant message found in the API response.'
    else:
        return f"Error: {response.status_code}, {response.text}"

# Worker function to process requests from the queue
def process_requests():
    while True:
        transcript_text, request_id, job_title = request_queue.get()
        if request_id is None:  # Stop signal
            break

        result = evaluate_interview(transcript_text, job_title)
        results[request_id] = result
        time.sleep(rate_limit_delay)  # Rate limiting delay
        request_queue.task_done()

# Start the worker thread
worker_thread = Thread(target=process_requests)
worker_thread.start()

# Function to generate a Word document of the evaluation feedback
def create_word_report(feedback):
    doc = Document()
    doc.add_heading('Interview Evaluation Feedback Report', level=1)

    for line in feedback.split('\n'):
        doc.add_paragraph(line)

    output = BytesIO()
    doc.save(output)
    output.seek(0)  # Move to the beginning of the BytesIO buffer
    return output

# Main app page
def main_app():
    st.title("Interview Insight Pro")
    st.write("Upload the interview transcript(s) to evaluate the candidate, interviewer, and responses using AI.")

    # Job title selection
    job_title_options = ["Software Engineer", "Data Analyst", "Product Manager", "Other"]
    selected_job_title = st.selectbox("Job Title", job_title_options)

    if selected_job_title == "Other":
        custom_job_title = st.text_input("Please specify the job title:")
        if custom_job_title:
            selected_job_title = custom_job_title

    # Experience level input
    experience_level = st.selectbox("Experience Level", ["Entry-Level", "Mid-Level", "Senior-Level"])

    # Multiple file upload
    uploaded_transcripts = st.file_uploader("Upload Interview Transcripts (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)

    if uploaded_transcripts:
        feedbacks = []  # To store feedbacks for each file

        for uploaded_transcript in uploaded_transcripts:
            transcript_text = extract_text_from_file(uploaded_transcript)

            # Input Validation: Check for valid transcript document upload
            if transcript_text is None:
                st.error(f"The uploaded transcript document {uploaded_transcript.name} is not valid. Please upload a valid PDF, DOCX, or TXT file.")
                continue

            # Sentiment analysis on the transcript text
            sentiment_score = analyze_sentiment(transcript_text)
            sentiment_emoji = "😊" if sentiment_score > 0 else "😐" if sentiment_score == 0 else "🥺"

            # Sentiment score and explanation
            st.write(f"**Overall Sentiment Score:** {sentiment_score:.3f} {sentiment_emoji}")
            
            # Confirmation message for valid uploads
            st.success(f"Transcript {uploaded_transcript.name} uploaded successfully! Ready for evaluation.")

            if st.button(f"Evaluate {uploaded_transcript.name}", key=f"evaluate_{uploaded_transcript.name}"):

                with st.spinner("Processing..."):
                    request_id = f"request_{time.time()}"  # Unique ID for this request
                    request_queue.put((transcript_text, request_id, selected_job_title))

                    # Display feedback in real-time
                    feedback_displayed = False  # Flag to track if feedback is displayed
                    while not feedback_displayed:
                        if len(results) > 0:
                            for request_id in list(results.keys()):
                                feedback = results.pop(request_id)  # Get the result and remove from the dictionary
                                st.text_area(f"Evaluation Feedback for {uploaded_transcript.name}", feedback, height=250, key=f"evaluation_feedback_{uploaded_transcript.name}")
                                feedback_displayed = True  # Stop the loop once feedback is displayed
                                break
                        time.sleep(1)  # Check every second

                    # Generate and allow downloading of the Word report
                    if feedback:  # Only create report if there's feedback
                        word_report = create_word_report(feedback)
                        st.download_button(f"Download Evaluation Report for {uploaded_transcript.name}", word_report, f"evaluation_report_{uploaded_transcript.name}.docx")

        # Explanation text after input and feedback output
        st.write("*Sentiment Score indicates the general tone of the responses, questions, and overall conversation.")
        st.write("**Negative Score** (close to -1) 🥺, **Neutral Score** (close to 0) 😐, **Positive Score** (close to +1) 😊")

    # Footer update with black and white text and reduced box size
    st.markdown(
        """
        <style>
        .footer {
            background-color: white;
            color: black;
            text-align: center;
            padding: 5px;
            font-size: 12px;
        }
        </style>
        """, 
        unsafe_allow_html=True
    )
    st.markdown(
        "<div class='footer'><p>Interview Insight Pro | Powered by Gen AI</p></div>", 
        unsafe_allow_html=True
    )

# Run the app
if __name__ == "__main__":
    main_app()
    request_queue.put((None, None, None))  # Stop the worker thread
    worker_thread.join()  # Wait for the worker thread to finish
